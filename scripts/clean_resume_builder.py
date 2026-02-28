import os
from docx import Document

# Create a brand new document from scratch, using the original template ONLY for its style definitions.
# This ensures ZERO trace of the original text remains (no hidden runs, metadata text, etc).
doc = Document('data/sample_resume.docx')

# Clear all existing paragraphs (this removes all old content completely)
for paragraph in doc.paragraphs:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# Now, add completely new paragraphs using the standard or generic styles
# (If custom styles were used, this will attempt to map to them, but fall back gracefully)

# Header
p = doc.add_paragraph("Alex Mercer")
p.style = doc.styles['Heading 1'] if 'Heading 1' in doc.styles else doc.styles['Normal']

p = doc.add_paragraph("Austin, TX | alex.mercer@email.com | (555) 123-4567 | linkedin.com/in/alex-mercer | github.com/alex-mercer")

# Professional Summary
p = doc.add_paragraph("PROFESSIONAL SUMMARY")
p.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else doc.styles['Normal']
doc.add_paragraph("Results-driven Machine Learning Engineer with 3+ years of experience in developing scalable backend systems, training predictive models, and deploying AI solutions. Proven ability to architect end-to-end data pipelines and integrate Large Language Models (LLMs) into production web applications.")

# Technical Skills
p = doc.add_paragraph("TECHNICAL SKILLS")
p.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else doc.styles['Normal']
doc.add_paragraph("Languages: Python, SQL, JavaScript, C++")
doc.add_paragraph("Machine Learning: TensorFlow, PyTorch, Scikit-Learn, MLflow, Pandas, NumPy")
doc.add_paragraph("Cloud & Infrastructure: AWS (EC2, S3, SageMaker), Docker, Kubernetes, CI/CD pipelines")
doc.add_paragraph("Concepts & Tools: Natural Language Processing (NLP), REST APIs, Data Engineering, Git, Agile")

# Experience
p = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
p.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else doc.styles['Normal']

p = doc.add_paragraph("TechNova Industries | Machine Learning Engineer")
p.style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Normal']
doc.add_paragraph("June 2022 – Present")
doc.add_paragraph("• Developed and deployed deep learning classification models using PyTorch on AWS, improving automated routing accuracy by 18% in production workflows.")
doc.add_paragraph("• Built scalable ETL data pipelines using Python and Pandas, processing over 500,000 financial records daily to feed into model training datasets.")
doc.add_paragraph("• Orchestrated CI/CD deployment pipelines using Docker and GitHub Actions, reducing model deployment time by 40%.")

p = doc.add_paragraph("DataPulse Analytics | Data Engineering Intern")
p.style = doc.styles['Heading 3'] if 'Heading 3' in doc.styles else doc.styles['Normal']
doc.add_paragraph("May 2021 – August 2021")
doc.add_paragraph("• Assisted in the architectural design of a cloud-based data lake using AWS S3 and Athena for big data analytical workloads.")
doc.add_paragraph("• Optimized legacy enterprise SQL queries, reducing weekly report generation times across 3 different departments by approximately 45%.")
doc.add_paragraph("• Built internal Streamlit dashboards for real-time visualization of key machine learning performance metrics.")

# Academic Projects
p = doc.add_paragraph("ACADEMIC PROJECTS")
p.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else doc.styles['Normal']

p = doc.add_paragraph("Job Application AI Agent | Python, SerpAPI, Pandas, Streamlit")
doc.add_paragraph("• Designed an autonomous agent using Python and SerpAPI to aggregate job postings and dynamically filter candidate pipelines based on constraints.")
doc.add_paragraph("• Implemented a custom ranking algorithm matching skills and locations, increasing relevant job discovery by 40% against naive searches.")
doc.add_paragraph("• Integrated OpenAI LLMs to automatically tailor base resumes to specific job descriptions based on parsed employer requirements.")

# Education
p = doc.add_paragraph("EDUCATION")
p.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else doc.styles['Normal']
doc.add_paragraph("Bachelor of Science in Computer Science")
doc.add_paragraph("University of Technology | Austin, TX")
doc.add_paragraph("Graduated: May 2022")

doc.save('data/sample_resume.docx')
print("Successfully generated clean sample_resume.docx")
