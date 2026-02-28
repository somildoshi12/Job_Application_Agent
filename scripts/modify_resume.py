from docx import Document
import re

doc = Document('data/Somil_Resume.docx')

# Replace exact strings in paragraphs to anonymize the resume template
replacements = {
    "Somil Doshi": "Alex Mercer",
    "doshi.somil12@gmail.com": "alex.mercer@email.com",
    "(704)-579-2423": "(555) 555-0123",
    "linkedin.com/in/somildoshi": "linkedin.com/in/alex-mercer",
    "github.com/somildoshi12": "github.com/alex-mercer"
}

for paragraph in doc.paragraphs:
    for old_text, new_text in replacements.items():
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

# Also check runs to preserve formatting during replacements where possible
for p in doc.paragraphs:
    for run in p.runs:
        for old_text, new_text in replacements.items():
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

doc.save('data/sample_resume.docx')
print("Saved modified resume to data/sample_resume.docx")
