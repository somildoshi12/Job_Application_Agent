from docx import Document
import re

doc = Document('data/sample_resume.docx')

# Complete replacement strategy to maintain format but change content
replacements = {
    # Skills
    "React, Next.js, Django, Node.js, Express.js": "TensorFlow, PyTorch, Scikit-Learn, MLflow",
    "PostgreSQL, MongoDB, Supabase, Redis, Qdrant, Chroma, VectorDB": "SQL, NoSQL, Redis, Vector Databases (Pinecone, Chroma)",
    "AWS, Docker, Vercel, Firebase, DigitalOcean, GCP": "AWS (EC2, S3, SageMaker), Docker, Kubernetes, CI/CD",
    "LLM Integration, LangChain, RAG, WebSockets, REST APIs, GraphQL, Cypress": "LLMs (OpenAI, Gemini), LangChain, Data Engineering, REST APIs",
    
    # Experience 1
    "AI SaaS Developer": "Machine Learning Engineer",
    "DevDoshi / Upwork": "TechNova Industries",
    "September 2024 – Present": "June 2022 – Present",
    "Engineered serverless Python backends using Supabase Edge Functions with Supabase Database Webhooks reducing latency by 35% compared to traditional continuous polling architectures.": "Developed and deployed deep learning classification models using PyTorch on AWS, improving automated routing accuracy by 18% in production.",
    "Integrated Qdrant and OpenAI embeddings to build a scalable Retrieval-Augmented Generation (RAG) system processing 10,000+ daily queries, improving context retrieval accuracy by 25%.": "Built scalable ETL data pipelines using Python and Pandas, processing over 500,000 financial records daily for model training datasets.",
    "Architected event-driven workflows with Supabase Realtime mapping database states enabling millisecond UI updates.": "Orchestrated CI/CD model deployment workflows using Docker and GitHub Actions, reducing model deployment time by 40%.",
    "Developed 15+ automated AI agents for healthcare and insurance clients leveraging Google Cloud infrastructure.": "Integrated Gemini LLM API into frontend React applications to automate customer service inquiries, reducing response times.",

    # Experience 2
    "Software Engineer Intern": "Data Engineering Intern",
    "Breezing": "DataPulse Analytics",
    "May 2024 – August 2024": "May 2021 – August 2021",
    "Developed an enterprise patient tracking portal utilizing React JS and Vite securing and managing records for 2,500+ active clinical patients.": "Assisted in the architectural design of a cloud-based data lake using AWS S3 and Athena for analytical workloads.",
    "Led migration from Firebase to AWS ecosystem incorporating S3 for image storage, reducing cloud overhead by 15%.": "Optimized legacy enterprise SQL queries, reducing weekly report generation times across 3 departments by approximately 45%.",
    "Built and optimized comprehensive REST APIs using Node.js and Express enabling bi-directional sync with mobile clients.": "Built internal Streamlit dashboards for real-time visualization of key machine learning performance metrics.",

    # Experience 3
    "Full Stack Freelancer": "Software Developer Intern",
    "September 2023 – April 2024": "June 2020 – August 2020",
    "Designed and launched 5 bespoke e-commerce platforms using Next.js processing $50,000+ in aggregate customer volume over 6 months.": "Developed internal automation scripts in Python, saving the engineering team 15 hours of manual work weekly.",
    "Engineered robust authentication layers via NextAuth.js and integrated Stripe webhooks handling 300+ secure transactions.": "Assisted senior developers in maintaining RESTful APIs using Flask and debugging production codebase issues.",
    
    # Projects
    "Projects": "Academic Projects",
    "Multifamily Real Estate Price Forecasting | Python, Scikit-Learn, Pandas, GIS, Streamlit": "Job Application AI Agent | Python, SerpAPI, Pandas, Streamlit, LLMs",
    "Built a revenue forecasting pipeline on 12,800+ multifamily property records, achieving 70.6% R² and 0.109 RMSE, reducing prediction error by 18% over baseline models.": "Designed an autonomous agent using Python and SerpAPI to aggregate job postings and dynamically filter candidate pipelines.",
    "Automated executive reporting using LLM-driven summaries and SHAP-based explainability, delivering interpretable investment insights and validating 15-minute city performance patterns via interactive dashboards.": "Implemented a custom ranking algorithm matching skills and locations, increasing relevant job discovery by 40% against naive searches.",
    "Predictive Maintenance Analytics for Oil Wells | Neural Networks, PostgreSQL, Docker, Streamlit": "Predictive Analytics Dashboard | Machine Learning, Streamlit",
    "Built data pipeline for drilling sensor data, improving maintenance lead time by 7 days through early issue detection.": "Built end-to-end data pipeline for drilling sensor data, visualizing results in a live Streamlit dashboard.",
    "Reduced downtime by 28% via data-driven maintenance scheduling across 5,000+ simulated well operations.": "Integrated OpenAI LLMs to automatically tailor base resumes to specific job descriptions based on parsed employer constraints.",
    
    # Remove extra project lines that spill onto page 2
    "Brain Tumor MRI Classification | 3D CNNs, PyTorch, Flask APIs, MongoDB Atlas, Docker": "",
    "Implemented a data pipeline for MRI inference, achieving 94% classification accuracy across multiple tumor subtypes.": "",
    "Logged and served predictions via Flask with MongoDB Atlas, reducing diagnostic reporting time by 30% overall.": ""
}

for paragraph in doc.paragraphs:
    for old_text, new_text in replacements.items():
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

# Also check runs
for p in doc.paragraphs:
    for run in p.runs:
        for old_text, new_text in replacements.items():
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

doc.save('data/sample_resume.docx')
print("Completed full content replacement on sample_resume.docx")
