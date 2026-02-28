from docx import Document
import re

def replace_all_info(doc_path, output_path):
    doc = Document(doc_path)
    
    # We will do a generic regex and string replacement across all runs
    # This targets the specific items seen in the user's screenshot
    
    replacements = {
        "Somil Doshi": "Alex Mercer",
        "doshi.somil12@gmail.com": "alex.mercer@email.com",
        "713.837.9474": "555.123.4567",
        "linkedin.com/in/somildoshi1202": "linkedin.com/in/alex-mercer",
        "github.com/somildoshi12": "github.com/alex-mercer",
        "Portfolio:somildoshi1202@gmail.com": "Portfolio: alexmercer.dev",
        
        # Education
        "University of Houston": "State University",
        "Aug 2024 – May 2026": "Aug 2020 – May 2024",
        "Master of Science in Engineering Data Science": "Bachelor of Science in Computer Science",
        "GPA: 3.89/4.0": "GPA: 3.8/4.0",
        "University of Mumbai": "City College",
        "Aug 2020 – May 2024": "Aug 2016 – May 2020",
        "Bachelor of Engineering in Information Technology": "Associate Degree in Software Engineering",
        "GPA: 3.84/4.0": "GPA: 3.7/4.0",
        
        # Experience 1
        "Research Assistant - Data Engineering and Analytics": "Machine Learning Engineer",
        "Mar 2025 – Dec 2026": "June 2022 – Present",
        "Designed and maintained end-to-end data pipelines for research datasets, supporting ingestion, transformation, and": "Developed and deployed deep learning classification models using PyTorch on AWS, improving automated",
        "structural storage, improving data availability and analysis readiness by 30%.": "routing accuracy by 18% in production workflows by designing scalable inference pipelines.",
        "Built and optimized SQL-based analytical data models using Databricks, reducing query latency by 22% and improving": "Built scalable ETL data pipelines using Python and Pandas, processing over 500,000 financial records daily",
        "analytical performance across research workflows.": "to feed into model training datasets and predictive analytics.",
        "Deployed and managed data services and storage systems using PostgreSQL, Docker, ensuring scalable, secure access": "Orchestrated CI/CD deployment pipelines using Docker and GitHub Actions, reducing model deployment",
        "to datasets and supporting multiple concurrent research projects.": "time by 40% while ensuring stable builds across environments.",
        
        # Experience 2
        "Golin Motors": "DataPulse Analytics",
        "Data Analyst Intern": "Data Engineering Intern",
        "Apr 2023 – Mar 2024": "May 2021 – August 2021",
        "Analyzed and consolidated 10K+ motor production, quality, and sales records using MongoDB and Databricks,": "Assisted in the architectural design of a cloud-based data lake using AWS S3 and Athena for analytical workloads,",
        "orchestrated with Airflow and standardized via dbt, improving forecasting reliability by 15%.": "improving data retrieval and forecasting reliability by 15%.",
        "Developed time-series analysis pipelines in Databricks, processing unstructured manufacturing and sales data, achieving": "Optimized legacy enterprise SQL queries, reducing weekly report generation times across 3 different departments",
        "92% accuracy and reducing prediction error by 18%.": "by 45% while achieving 92% historical accuracy.",
        "Designed Power BI dashboards with DAX to track motor performance KPIs, production throughput, and sales sea-": "Built internal Streamlit dashboards for real-time visualization of key machine learning performance KPIs",
        "sonality, helping operations and commercial teams identify anomalies and plan capacity effectively.": "helping operations and commercial teams identify anomalies effectively.",
        
        # Experience 3
        "TCR Innovations": "TechNova Industries",
        "Data Engineer Intern": "Software Developer Intern",
        "Jan 2022 – Apr 2023": "June 2020 – August 2020",
        "Introduced ETL pipelines in Python with relational database integration, managing 40K+ drilling records and enforcing": "Developed internal automation scripts in Python, saving the engineering team 15 hours of manual work weekly",
        "schema rules, reducing refresh cycles by 34% and improving reliability of downstream workloads.": "and enforcing strict schema rules, improving reliability of downstream workloads.",
        "Deployed predictive models supporting operational analytics, linked to React.js interfaces, sustaining 89% accuracy": "Assisted senior developers in maintaining RESTful APIs using Flask, linked to React.js interfaces, sustaining",
        "while reducing manual intervention.": "99.9% uptime while reducing manual intervention.",
        "Automated deployments with Docker, ensuring stable builds across staging and production environments.": "Automated deployments with Docker, ensuring stable builds across staging and production environments.",
        
        # Projects 1
        "Projects": "Academic Projects",
        "RevPAR Intelligence Platform - BroadVail Capital | Python, Streamlit, OpenAI API, XGBoost, LightGBM, SHAP": "Job Application AI Agent | Python, SerpAPI, Pandas, Streamlit, LLMs",
        "Built a revenue forecasting pipeline on 12,800+ multifamily property records, achieving 70.6% R² and 0.109 RMSE, reducing prediction error by 18% over baseline models.": "Designed an autonomous agent using Python and SerpAPI to aggregate job postings and dynamically filter candidate pipelines, reducing manual search time by 80%.",
        "Automated executive reporting using LLM-driven summaries and SHAP-based explainability, delivering interpretable investment insights and validating 15-minute city performance patterns via interactive dashboards.": "Implemented a custom ranking algorithm matching skills and locations, increasing relevant job discovery by 40% against naive searches, validating matches via interactive dashboards.",
        
        # Projects 2 & 3
        "Predictive Maintenance Analytics for Oil Wells | Neural Networks, PostgreSQL, Docker, Streamlit": "Predictive Analytics Dashboard | Machine Learning, Streamlit",
        "Built data pipeline for drilling sensor data, improving maintenance lead time by 7 days through early issue detection.": "Built end-to-end data pipeline for user behavior sensor data, improving retention lead time by 7 days through early detection.",
        "Reduced downtime by 28% via data-driven maintenance scheduling across 5,000+ simulated well operations.": "Reduced churn by 28% via data-driven intervention scheduling across 5,000+ simulated user operations.",
        
        "Brain Tumor MRI Classification | 3D CNNs, PyTorch, Flask APIs, MongoDB Atlas, Docker": "Customer Churn Prediction Model | XGBoost, Scikit-Learn",
        "Implemented a data pipeline for MRI inference, achieving 94% classification accuracy across multiple tumor subtypes.": "Developed a predictive model identifying high-risk churn customers with 89% AUC across multiple user subtypes.",
        "Logged and served predictions via Flask with MongoDB Atlas, reducing diagnostic reporting time by 30% overall.": "Deployed inference API that integrated with CRM system to automate retention email flagging, saving 30% reporting time."
    }

    # Pass 1: Replace inside runs (preserves exact run formatting if text is contained within one run)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for old_text, new_text in replacements.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

    # Pass 2: Deal with split runs
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        needs_replacement = False
        for old_text in replacements:
            if old_text in full_text:
                needs_replacement = True
                break
                
        if needs_replacement:
            # Reconstruct the string
            new_full_text = full_text
            for old_text, new_text in replacements.items():
                new_full_text = new_full_text.replace(old_text, new_text)
            
            # Apply to first run and clear others so formatting of first word dictates the block
            if paragraph.runs:
                paragraph.runs[0].text = new_full_text
                for r in paragraph.runs[1:]:
                    r.text = ""
                    
    # Also iterate URLs/Hyperlinks which are stored differently in python-docx
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype:
            if "linkedin.com/in/somildoshi1202" in rel._target:
                rel._target = "https://linkedin.com/in/alex-mercer"
            elif "github.com/somildoshi12" in rel._target:
                rel._target = "https://github.com/alex-mercer"
            elif "somildoshi12o2@gmail.com" in rel._target:
                 rel._target = "mailto:alex.mercer@email.com"

    doc.save(output_path)

if __name__ == "__main__":
    replace_all_info("data/Somil_Resume.docx", "data/sample_resume.docx")
    print("Complete Deep Replacement.")
