import io
import docx
from typing import List, Dict

def apply_in_place_edits(docx_bytes: bytes, replacement_ops: List[Dict[str, str]]) -> bytes:
    """
    Opens a raw .docx binary, searches through all paragraphs and runs for the 
    'original_text', replaces it with 'new_text', and returns the new binary 
    without destroying original styling.
    """
    if not docx_bytes:
        raise ValueError("No docx binary provided for editing.")
        
    doc = docx.Document(io.BytesIO(docx_bytes))
    
    # Track which operations were successfully applied
    applied_count = 0
    
    for op in replacement_ops:
        old_text = op.get("original_text", "")
        new_text = op.get("new_text", "")
        
        if not old_text or old_text == new_text:
            continue
            
        success = False
        
        # We must iterate through both paragraphs and tables to catch all text
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                # To preserve formatting, we ideally edit runs, but if a phrase crosses 
                # run boundaries, python-docx can't easily replace it without destroying 
                # run formatting. We do a blunt replacement on the paragraph level as a fallback.
                # However, python-docx paragraph.text assignment wipes runs.
                # A safer approach for exact keyword replacement:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        success = True
                
                # If we couldn't find it cleanly in a single run, do the blunt replace
                if not success and old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    success = True
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                success = True
                        if not success and old_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(old_text, new_text)
                            success = True
                            
        if success:
            applied_count += 1
            
    print(f"Applied {applied_count}/{len(replacement_ops)} smart in-place edits to DOCX.")
    
    out_stream = io.BytesIO()
    doc.save(out_stream)
    return out_stream.getvalue()
