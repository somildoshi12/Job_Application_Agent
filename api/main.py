from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
import os
import sys
import json
import io
import re

import docx
import PyPDF2

from fastapi.middleware.cors import CORSMiddleware

# Add parent directory to path so we can import the src logic
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from src.pipeline import run_pipeline

app = FastAPI(title="Job Application Agent API")

# Add CORS middleware to allow React to communicate with FastAPI
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # Allow all origins for local dev
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/api/run-pipeline")
async def execute_pipeline(
    query: str = Form(...),
    location: str = Form(...),
    user_skills: str = Form(...),
    target_top_k: int = Form(3),
    resume_file: UploadFile = File(...)
):
    try:
        content = await resume_file.read()
        filename = resume_file.filename.lower()
        
        base_resume_text = ""
        if filename.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                base_resume_text += page.extract_text() + "\n"
        elif filename.endswith(".docx") or filename.endswith(".doc"):
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                base_resume_text += para.text + "\n"
        elif filename.endswith(".txt"):
            base_resume_text = content.decode('utf-8')
        else:
            raise HTTPException(400, "Unsupported file type. Please upload PDF, DOCX, or TXT.")

        skills_list = [s.strip() for s in user_skills.split(",")]

        # Run the core logic
        results = run_pipeline(
            query=query,
            location=location,
            user_skills=skills_list,
            base_resume_text=base_resume_text,
            base_resume_bytes=content if filename.endswith(('.docx', '.doc')) else None,
            target_top_k=target_top_k
        )
        
        # Load the trace log from disk so we can return it sequentially
        trace_data = []
        if os.path.exists(results["trace_file"]):
            with open(results["trace_file"], "r") as f:
                trace_data = json.load(f)
                
        # We need to serialize the Pydantic Job objects into dicts for FastAPI to build the JSON response correctly
        ranked_jobs = [j.model_dump() for j in results["ranked_jobs"]]
                
        return {
            "status": "success",
            "ranked_jobs": ranked_jobs,
            "tailored_applications": results["tailored_results"],
            "trace_log": trace_data,
            "trace_file": results["trace_file"]
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class DocxFromTextRequest(BaseModel):
    text: str
    company_name: str

@app.post("/api/download-docx")
def download_cover_letter_docx(req: DocxFromTextRequest):
    """Generates a nicely formatted .docx from markdown text (used for Cover Letters)."""
    doc = docx.Document()
    doc.add_heading(f"Cover Letter – {req.company_name}", 0)

    for line in req.text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            is_bullet = line.startswith("- ") or line.startswith("* ")
            if is_bullet:
                p = doc.add_paragraph(style='List Bullet')
                text_content = line[2:]
            else:
                p = doc.add_paragraph()
                text_content = line
            parts = text_content.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    filename = f"{req.company_name.replace(' ', '_')}_Cover_Letter.docx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


class DownloadTailoredDocxRequest(BaseModel):
    b64_bytes: str
    company_name: str

@app.post("/api/download-tailored-docx")
def download_tailored_docx(req: DownloadTailoredDocxRequest):
    """
    Decodes the dynamically-modified base64 DOCX binary returned by the Agent pipeline
    and prompts the user's browser to download it as a native Word Document.
    """
    import base64
    raw_bytes = base64.b64decode(req.b64_bytes)
    
    file_stream = io.BytesIO(raw_bytes)
    filename = f"{req.company_name.replace(' ', '_')}_Tailored_Resume.docx"
    
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/")
def health_check():
    return {"status": "Agent API is live."}
